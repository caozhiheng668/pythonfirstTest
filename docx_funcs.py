import json
from docx import Document

POOL_SIZE = 15
LEN_LIMIT = 100

BREAK_STARTS = [
    "\u0f14"

]

BREAK_ENDS = [
    "\u0f14"
]

MAX_SPACE = 1

BREAK_SEQUENCE = []

for NUM_SPACE in range(MAX_SPACE+1):
    for BREAK_START in BREAK_STARTS:
        for BREAK_END in BREAK_ENDS:
            SPACE_STRING = " " * NUM_SPACE
            BREAK_SEQUENCE.append(f"{BREAK_START}{SPACE_STRING}{BREAK_END}")

def preprocess(text):
    """Before breaking paragraphs"""
    return text.replace("\n", "").replace("\r", "")

def postprocess(text):
    """After breaking paragraphs"""
    return text.replace(" ", "\u00A0") # Replace with non-breaking space for better layout

def break_text(text):
    """Break a full chunk of text without line breaks into paragraphs (splitting after ||  ||)"""
    texts = []
    i = 0
    next_i = 0
    end = len(text)
    while next_i < end:
        next_i = end
        for sequence in BREAK_SEQUENCE:
            found_i = text.find(sequence, i)
            if found_i >= 0:
                next_i = min(next_i, found_i + len(sequence))
        texts.append(text[i:next_i])
        i = next_i
    return texts

def create_docx(kb_name, pb_name):

    # CHANGE THIS
    # load data from file
    with open(f"doc/{kb_name}/{pb_name}.json", "r", encoding="utf-8") as text_file:
        text_obj = json.load(text_file)

    try:
        para_count = 0

        document = Document("docx_template.docx") # Open the template, which has heading and body styles configured
        titles = []
        paragraph_buff = []
        # CHANGE THIS
        for text_item in text_obj:
            if text_item["t"] == "bio":
                # Heading
                title = text_item["data"]["tname"]
                titles.append(title)
                if paragraph_buff:
                    for paragraph_text in break_text("".join(paragraph_buff)):
                        document.add_paragraph(postprocess(paragraph_text))
                        para_count+=1
                    paragraph_buff = []
                document.add_heading(title, level=1)
            elif text_item["t"] == "text":
                # Body
                text = text_item["text"]
                paragraph_buff.append(preprocess(text))

        if paragraph_buff:
            for paragraph_text in break_text("".join(paragraph_buff)):
                document.add_paragraph(postprocess(paragraph_text))
                para_count+=1
        
        # Process title
        full_title = "+".join(titles)
        if len(full_title) > LEN_LIMIT:
            full_title=full_title[:LEN_LIMIT-3]+"..."

        document.save(f"docx/{kb_name}/{pb_name}-{full_title}.docx")
    except ValueError:
        id = text_item["id"]
        with open(f"err/docx-{kb_name}-{id}.txt", "w+", encoding="utf-8") as err_file:
            err_file.write(f"Value Error for text: {text}\n")

    print(f"{kb_name}/{pb_name} done")

