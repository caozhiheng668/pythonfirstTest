import json
import jsonpath
from lxml import etree
import json
from docx import Document
#obj = json.load(open('dbfz_TD1_1-1-1a.json','r',encoding='utf-8'))
#print(obj)
# xpath解析本地文件
#tree = etree.parse('dbfz_TD1_1-1-1a.json')
#print(tree)

POOL_SIZE = 15
LEN_LIMIT = 100

BREAK_STARTS = [
    "\u0f14"

]

BREAK_ENDS = [
    "\u0f14"
]

MAX_SPACE = 1

BREAK_SEQUENCE = []

for NUM_SPACE in range(MAX_SPACE+1):
    for BREAK_START in BREAK_STARTS:
        for BREAK_END in BREAK_ENDS:
            SPACE_STRING = " " * NUM_SPACE
            BREAK_SEQUENCE.append(f"{BREAK_START}{SPACE_STRING}{BREAK_END}")

def preprocess(text):
    """Before breaking paragraphs"""
    return text.replace("\n", "").replace("\r", "")

def postprocess(text):
    """After breaking paragraphs"""
    return text.replace(" ", "\u00A0") # Replace with non-breaking space for better layout

def break_text(text):
    """Break a full chunk of text without line breaks into paragraphs (splitting after ||  ||)"""
    texts = []
    i = 0
    next_i = 0
    end = len(text)
    while next_i < end:
        next_i = end
        for sequence in BREAK_SEQUENCE:
            found_i = text.find(sequence, i)
            if found_i >= 0:
                next_i = min(next_i, found_i + len(sequence))
        texts.append(text[i:next_i])
        i = next_i
    return texts

def retAddress_List():
    # xpath解析本地文件
    tree = etree.parse('idx.html')
    # 查询id的值以l开头的li标签
    li_list = tree.xpath('//li/ul/li/@aria-labelledby')
    # 判断列表的长度
    # print(li_list)
    # print(len(li_list))
    address_List = []
    for i in li_list:
        x = i.split("_anchor")
        zuHe = x[0].split("TD")[1]
        #  print(zuHe)
        chai = zuHe.split("-")  # chai[0]=11 chai[1]=1 chai[2]=1a
        ss = chai[0] + "-" + chai[1] + "-" + "1a"
        tdLen = len(chai[0])  # 11  101  2979  72218  257758
        if tdLen == 2:
            diYi = "TD" + chai[0][0] + "," + chai[0][1] + "-" + chai[1] + "-" + "1a"
            address_List.append(diYi)
        #    print(diYi)
        elif tdLen == 3:
            diYi = "TD" + str(chai[0][0]) + str(chai[0][1]) + "," + str(chai[0][2]) + "-" + chai[1] + "-" + "1a"
            address_List.append(diYi)
        #   print(diYi)
        elif tdLen == 4 or tdLen == 5:
            if tdLen == 4:
                diYi = "TD" + str(chai[0][0]) + str(chai[0][1]) + str(chai[0][2]) + "," + str(chai[0][3]) + "-" + chai[
                    1] + "-" + "1a"
                address_List.append(diYi)
            #   print(diYi)
            else:
                diYi = "TD" + str(chai[0][0]) + str(chai[0][1]) + str(chai[0][2]) + "," + str(chai[0][3]) + str(
                    chai[0][4]) + "-" + chai[
                           1] + "-" + "1a"
                address_List.append(diYi)
            #   print(diYi)

        elif tdLen == 6:
            diYi = "TD" + str(chai[0][0]) + str(chai[0][1]) + str(chai[0][2]) + str(chai[0][3]) + "," + str(
                chai[0][4]) + str(chai[0][5]) + "-" + chai[1] + "-" + "1a"
            address_List.append(diYi)

    return address_List


def logical_Processor(fileName,fileStrtext,catalog):
    windex = 1
    with open(fileName,encoding='utf-8') as f:
      data = json.load(f)
    #  print(type(data))
      reaptName = ""
      headerLVNext = ''
      LVnextIndex = 0
      kaiGuan = 0
      headContent = 0
      dataCount = len(data)
      if dataCount == 1:
          return []
      step = 0;
      for ds in data:
            step += 1
            if step == 1:
                continue
            pbName = ds.get('pbName')
            text = ds.get('text')
            headerLV = ds.get('headerLV')
            tree = etree.HTML(text)
            #span中class的属性值
            sclass = tree.xpath('//body//span/@class')[0]

            dataLVS = tree.xpath('//body//span/@data-lv')
            dataLV = ""
            if dataLVS:
                dataLV = dataLVS[0]




            if sclass == "head" :
                kaiGuan = 1
            if kaiGuan == 1:
                if fileStrtext != '' and headContent == 0:
                    kk = preprocess(fileStrtext)
                    k2 = postprocess(kk)
                    document.add_paragraph(k2)
                    print("===========" + k2)
                    fileStrtext = ''
                    headerLVNext = ''
                    LVnextIndex = 0
            if sclass == "head":
                print(sclass)
                if reaptName == pbName :
                    print()
                else:
                    print(pbName)
                    reaptName = pbName
                data_t = tree.xpath('//body//span/@data-t')[0]
                print(data_t)
                document.add_heading(data_t, level=1)

                # 添加目录
                if dataLV == "1":
                    catalog.append(dataLV + "-" + data_t)
                elif dataLV == "1.1":
                    catalog.append(dataLV + "-" + data_t)
                elif dataLV == "1.2":
                    catalog.append(dataLV + "-" + data_t)
                elif dataLV == "1.3":
                    catalog.append(dataLV + "-" + data_t)
                elif dataLV == "1.4":
                    catalog.append(dataLV + "-" + data_t)
                elif dataLV == "1.5":
                    catalog.append(dataLV + "-" + data_t)
                elif dataLV == "1.6":
                    catalog.append(dataLV + "-" + data_t)
                elif dataLV == "1.7":
                    catalog.append(dataLV + "-" + data_t)
                elif dataLV == "1.8":
                    catalog.append(dataLV + "-" + data_t)
                elif dataLV == "1.9":
                    catalog.append(dataLV + "-" + data_t)


                if windex == 1:
                    wname = data_t
                    windex+=1
                content = tree.xpath('//body/text()')
                if not content:
                    print()
                else:
                    print(content)
                    #document.add_paragraph(content[0])
                    headContent = 1
                    for ct in content:
                        fileStrtext += ct

            else:
                headContent = 0
                kaiGuan = 0
                reaptName = ""
                LVnextIndex += 1
                if LVnextIndex == 1:
                    headerLVNext = headerLV
                if headerLVNext == headerLV:
                    content = tree.xpath('//body/text()')
                    #判断结合是否是空
                    if not content:
                        print()
                    else:
                    #    document.add_paragraph(content[0])
                        print(pbName)
                        print(content)
                        for ct in content:
                            fileStrtext += ct

    kk = preprocess(fileStrtext)
    k2 = postprocess(kk)
    document.add_paragraph(k2)
    fileStrtext = ''
    headerLVNext = ''
    LVnextIndex = 0
    print("===========" + k2)
    return data







addressList = retAddress_List()
print(addressList)
print(len(addressList))

for ad in addressList:
    document = Document("docx_template.docx")
    tp = ad.split(",")
    tdName = tp[0]
    aa = tdName
    bb = aa.split("TD")[1]
    if int(bb) < 2653:
        continue
    pageName = tp[1]
    fileName = "dbfz_"+str(tdName)+ "_"+pageName+".json"

    if fileName == "dbfz_TD2654_65-1a-1a.json":
        fileName = "dbfz_TD2654_65-32a.json"

    if fileName == "dbfz_TD2654_66-1a-1a.json":
        fileName = "dbfz_TD2654_66-9b.json"

    if fileName == "dbfz_TD2654_67-1a-1a.json":
        fileName = "dbfz_TD2654_67-23b.json"
    wordPre = tp[1].split("1a")[0]
    print(fileName)
    print(wordPre)
    fileStrtext = ''
    wname = ""
    windex = 1
    catalog = []
    with open(fileName,encoding='utf-8') as f:
      data = json.load(f)
    #  print(type(data))
      reaptName = ""
      headerLVNext = ''
      LVnextIndex = 0
      kaiGuan = 0
      headContent = 0
      for ds in data:
            pbName = ds.get('pbName')
            text = ds.get('text')
            headerLV = ds.get('headerLV')
            tree = etree.HTML(text)
            #span中class的属性值
            sclass = tree.xpath('//body//span/@class')[0]

            dataLVS = tree.xpath('//body//span/@data-lv')
            dataLV = ""
            if dataLVS :
                dataLV = dataLVS[0]

            if sclass == "head" :
                kaiGuan = 1
            if kaiGuan == 1:
                if fileStrtext != '' and headContent == 0:
                    kk = preprocess(fileStrtext)
                    k2 = postprocess(kk)
                    document.add_paragraph(k2)
                    print("===========" + k2)
                    fileStrtext = ''
                    headerLVNext = ''
                    LVnextIndex = 0
            if sclass == "head":

                print(sclass)
                if reaptName == pbName :
                    print()
                else:
                    print(pbName)
                    reaptName = pbName
                data_t = tree.xpath('//body//span/@data-t')[0]
                print(data_t)
                document.add_heading(data_t, level=1)

                # 添加目录
                if dataLV == "1":
                    catalog.append(dataLV +"-"+ data_t)
                elif dataLV == "1.1":
                    catalog.append(dataLV +"-"+ data_t)
                elif dataLV == "1.2":
                    catalog.append(dataLV +"-"+ data_t)
                elif dataLV == "1.3":
                    catalog.append(dataLV +"-"+ data_t)
                elif dataLV == "1.4":
                    catalog.append(dataLV +"-"+ data_t)
                elif dataLV == "1.5":
                    catalog.append(dataLV +"-"+ data_t)
                elif dataLV == "1.6":
                    catalog.append(dataLV +"-"+ data_t)
                elif dataLV == "1.7":
                    catalog.append(dataLV +"-"+ data_t)
                elif dataLV == "1.8":
                    catalog.append(dataLV +"-"+ data_t)
                elif dataLV == "1.9":
                    catalog.append(dataLV +"-"+ data_t)


                if windex == 1:
                    wname = data_t
                    windex+=1
                content = tree.xpath('//body/text()')
                if not content:
                    print()
                else:
                    print(content)
                    #document.add_paragraph(content[0])
                    headContent = 1
                    for ct in content:
                        fileStrtext += ct

            else:
                headContent = 0
                kaiGuan = 0
                reaptName = ""
                LVnextIndex += 1
                if LVnextIndex == 1:
                    headerLVNext = headerLV
                if headerLVNext == headerLV:
                    content = tree.xpath('//body/text()')
                    #判断结合是否是空
                    if not content:
                        print()
                    else:
                    #    document.add_paragraph(content[0])
                        print(pbName)
                        print(content)
                        for ct in content:
                            fileStrtext += ct

    kk = preprocess(fileStrtext)
    k2 = postprocess(kk)
    document.add_paragraph(k2)
    fileStrtext = ''
    headerLVNext = ''
    LVnextIndex = 0
    print("===========" + k2)


    while True:
        if len(data) > 0:
            lens = len(data)
            if lens == 1 or lens == 2 or lens == 3:
                break
                print("正在写入下一个TD")
            # 最后一个json对象
            lastE = data[lens - 1]
            # 下一页的页码
            next_pbName = lastE.get("pbName")
            # 请求文件对象的定制
            next_fileName = "dbfz_" + str(tdName) + "_" + next_pbName + ".json"
            # 获取文件内容
            next_data = logical_Processor(next_fileName,fileStrtext,catalog)
            data = next_data

        else:
            print("先暂时结束")
            break
    #最后写入目录
    document.add_heading("目录")
    for cata in catalog:
        document.add_heading(cata)

    finalWordName =  wordPre + wname
    if len(finalWordName) > LEN_LIMIT:
        wordName = finalWordName[:LEN_LIMIT - 3] + "..."
        finalWordName = wordName

    document.save(finalWordName+'.docx')
